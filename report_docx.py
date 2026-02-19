# report_docx.py
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

try:
    from PIL import Image  # type: ignore
except Exception:  # pillow optional at runtime
    Image = None  # type: ignore


def _fmt_dt(dt: datetime) -> str:
    return dt.strftime("%d.%m.%Y, %H:%M:%S")


def _safe_text(x: Any) -> str:
    s = "" if x is None else str(x)
    s = s.strip()
    return s if s else "—"


def _clear_cell_keep_format(cell) -> None:
    """
    Очищает текст в ячейке, сохраняя форматирование (tcPr / pPr),
    оставляя один первый параграф.
    """
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if cell.paragraphs:
        p0 = cell.paragraphs[0]

        # удаляем лишние параграфы
        for p in cell.paragraphs[1:]:
            cell._tc.remove(p._p)

        # очищаем текст в runs первого параграфа
        if p0.runs:
            for r in p0.runs:
                r.text = ""
        else:
            p0.add_run("")
    else:
        cell.add_paragraph("")

    if not cell.paragraphs:
        cell.add_paragraph("")


def _remove_paragraph_numbering(paragraph) -> None:
    """
    Убирает нумерацию списка (w:numPr) у параграфа, если она есть.
    Нужно, чтобы Word не добавлял автонумерацию поверх нашего номера,
    иначе получится "11", "22", "33" и т.п.
    """
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def _to_png_bytes(image_bytes: bytes) -> bytes | None:
    """
    Конвертирует исходные bytes (jpeg/png/webp/...) в PNG bytes через Pillow.
    Возвращает None, если Pillow недоступен или данные не распознаны.
    """
    if not image_bytes:
        return None
    if Image is None:
        return None
    try:
        im = Image.open(BytesIO(image_bytes))
        im.load()

        # нормализуем режим для сохранения в PNG
        if im.mode not in ("RGB", "RGBA"):
            if "A" in im.mode:
                im = im.convert("RGBA")
            else:
                im = im.convert("RGB")

        out = BytesIO()
        im.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def _safe_add_picture(run, image_bytes: bytes, width: Inches) -> bool:
    """
    Пытается вставить картинку:
      1) как есть
      2) если не распознано — конвертит в PNG через Pillow
    НИКОГДА не кидает исключение — возвращает True/False.
    """
    # 1) напрямую
    try:
        run.add_picture(BytesIO(image_bytes), width=width)
        return True
    except Exception:
        pass

    # 2) через PNG
    png = _to_png_bytes(image_bytes)
    if not png:
        return False

    try:
        run.add_picture(BytesIO(png), width=width)
        return True
    except Exception:
        return False


def build_defect_report_docx_bytes(
    *,
    template_path: str | Path,
    created_at: datetime,
    rows: list[dict[str, Any]],
) -> bytes:
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    # 1) Обновляем строку "Сформировано: ..."
    stamp = _fmt_dt(created_at)
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt.startswith("Сформировано:"):
            for r in p.runs:
                r.text = ""
            run = p.add_run(f"Сформировано: {stamp}")
            run.italic = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2) Первая таблица
    if not doc.tables:
        raise RuntimeError("Template has no tables")

    table = doc.tables[0]
    if len(table.rows) < 2:
        raise RuntimeError("Template table must contain header + at least one sample row")

    sample_row = table.rows[1]

    # 3) Удаляем строки данных (оставляем заголовок)
    for i in range(len(table.rows) - 1, 0, -1):
        table._tbl.remove(table.rows[i]._tr)

    def add_row_like_sample():
        new_tr = deepcopy(sample_row._tr)
        table._tbl.append(new_tr)
        return table.rows[-1]

    # 4) Заполняем
    # ВАЖНО: "№ п/п" делаем 1..N
    for i, item in enumerate(rows, start=1):
        r = add_row_like_sample()

        for c in r.cells:
            _clear_cell_keep_format(c)

        idx = i
        place = _safe_text(item.get("place"))
        category = _safe_text(item.get("category"))
        recommendation = _safe_text(item.get("recommendation"))
        desc_lines: list[str] = item.get("description_lines") or ["—"]
        if not isinstance(desc_lines, list) or not desc_lines:
            desc_lines = ["—"]

        # № (центр) + FIX: убираем автонумерацию списка из шаблона
        c0 = r.cells[0]
        p0 = c0.paragraphs[0]
        _remove_paragraph_numbering(p0)

        # часто в шаблоне стоит стиль "List Paragraph" -> нормализуем в "Normal"
        try:
            p0.style = doc.styles["Normal"]
        except Exception:
            pass

        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.add_run(str(idx))

        # Расположение (центр)
        c1 = r.cells[1]
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(place)

        # Фото (центр) + FIX: webp/битые bytes не должны ронять генерацию
        img_cell = r.cells[2]
        img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pimg = img_cell.paragraphs[0]
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER

        image_bytes: bytes | None = item.get("image_bytes")
        if isinstance(image_bytes, bytearray):
            image_bytes = bytes(image_bytes)
        if not isinstance(image_bytes, (bytes, type(None))):
            image_bytes = None

        if image_bytes:
            run = pimg.add_run()
            ok = _safe_add_picture(run, image_bytes, width=Inches(2.1))
            if not ok:
                pimg.add_run("(изображение недоступно)")
        else:
            pimg.add_run("(нет фото)")

        # Описание
        desc_cell = r.cells[3]
        desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pdesc0 = desc_cell.paragraphs[0]
        pdesc0.add_run(str(desc_lines[0]))

        for line in desc_lines[1:]:
            p = desc_cell.add_paragraph(str(line))
            p.style = pdesc0.style
            p.alignment = pdesc0.alignment

        # Категория (центр)
        c4 = r.cells[4]
        p4 = c4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run(category)

        # Рекомендация (центр)
        c5 = r.cells[5]
        p5 = c5.paragraphs[0]
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.add_run(recommendation)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
